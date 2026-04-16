from __future__ import annotations
import re
from pathlib import Path
from typing import Optional

from litmap.zotero import get_all_items, Item, ZOTERO_DB

# DOI pattern: matches 10.XXXX/anything
_DOI_RE = re.compile(r'\b(10\.\d{4,}/\S+?)(?:[,;)\s]|$)')


def extract_dois(text: str) -> list[str]:
    return [m.group(1).rstrip(".") for m in _DOI_RE.finditer(text)]


# --- BibTeX parser -----------------------------------------------------------

_ENTRY_RE = re.compile(r'@\w+\s*\{[^@]+', re.DOTALL)
_FIELD_RE = re.compile(r'(\w+)\s*=\s*[\{"](.*?)[\}"]', re.DOTALL)


def parse_bib_text(text: str) -> list[dict[str, str]]:
    """Parse a BibTeX string into a list of field dicts."""
    entries = []
    for m in _ENTRY_RE.finditer(text):
        block = m.group(0)
        fields = {k.lower(): v.strip() for k, v in _FIELD_RE.findall(block)}
        if fields:
            entries.append(fields)
    return entries


def _parse_bib_file(path: Path) -> list[dict[str, str]]:
    return parse_bib_text(path.read_text(encoding="utf-8", errors="replace"))


def _parse_tex_file(path: Path) -> list[dict[str, str]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    # Extract \bibitem{key} ... content
    entries = []
    for m in re.finditer(r'\\bibitem(?:\[.*?\])?\{([^}]+)\}(.*?)(?=\\bibitem|\\end\{thebibliography\})', text, re.DOTALL):
        content = m.group(2).strip()
        dois = extract_dois(content)
        entry = {"key": m.group(1), "raw": content}
        if dois:
            entry["doi"] = dois[0]
        entries.append(entry)
    return entries


def _parse_pdf(path: Path) -> list[dict[str, str]]:
    import fitz  # PyMuPDF
    doc = fitz.open(str(path))
    full_text = "\n".join(page.get_text() for page in doc)
    doc.close()
    return _extract_refs_from_text(full_text)


def _parse_docx(path: Path) -> list[dict[str, str]]:
    from docx import Document
    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    return _extract_refs_from_text(full_text)


def _extract_refs_from_text(text: str) -> list[dict[str, str]]:
    """Find References section in plain text and extract DOIs."""
    # Locate references section
    ref_match = re.search(
        r'(?m)^(References|Bibliography|Works Cited|Literature Cited)\s*$', text
    )
    if ref_match:
        text = text[ref_match.start():]
    dois = extract_dois(text)
    return [{"doi": d} for d in dict.fromkeys(dois)]  # deduplicate, preserve order


def parse_bibliography(path: Path) -> list[dict[str, str]]:
    """Parse bibliography from PDF, DOCX, .bib, or .tex file."""
    suffix = path.suffix.lower()
    if suffix == ".bib":
        return _parse_bib_file(path)
    if suffix == ".tex":
        return _parse_tex_file(path)
    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix in (".docx",):
        return _parse_docx(path)
    raise ValueError(f"Unsupported manuscript format: {suffix}")


def extract_manuscript_text(path: Path) -> str:
    """Extract title + abstract (or first ~500 words) for embedding."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        import fitz
        doc = fitz.open(str(path))
        text = doc[0].get_text() if len(doc) > 0 else ""
        doc.close()
    elif suffix == ".docx":
        from docx import Document
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
    elif suffix == ".tex":
        text = path.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported format: {suffix}")

    # Try to extract abstract section
    abs_match = re.search(
        r'(?i)(abstract|summary)[:\s]*\n?(.*?)(?=\n\n|\Z)', text, re.DOTALL
    )
    if abs_match:
        return abs_match.group(2).strip()[:2000]
    # Fall back to first ~500 words
    words = text.split()
    return " ".join(words[:500])


# --- Zotero matching ---------------------------------------------------------

def match_items_to_zotero(
    entries: list[dict[str, str]],
    zotero_db: Path = ZOTERO_DB,
) -> tuple[list[Item], list[dict[str, str]]]:
    """Match bibliography entries to Zotero items by DOI, then title.

    Returns (matched_items, unmatched_entries).
    """
    from difflib import SequenceMatcher

    all_items = get_all_items(zotero_db)
    doi_index = {i.doi.lower(): i for i in all_items if i.doi}
    title_index = {i.title.lower(): i for i in all_items}

    matched, unmatched = [], []
    seen_keys: set[str] = set()

    for entry in entries:
        doi = entry.get("doi", "").lower().strip()
        title = entry.get("title", "").lower().strip()

        item = None
        if doi and doi in doi_index:
            item = doi_index[doi]
        elif title:
            # Fuzzy title match with threshold 0.85
            best_score, best_item = 0.0, None
            for t, candidate in title_index.items():
                score = SequenceMatcher(None, title, t).ratio()
                if score > best_score:
                    best_score, best_item = score, candidate
            if best_score >= 0.85:
                item = best_item

        if item and item.key not in seen_keys:
            seen_keys.add(item.key)
            matched.append(item)
        else:
            unmatched.append(entry)

    return matched, unmatched
