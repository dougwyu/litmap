#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "python-docx>=1.0.0",
# ]
# ///
"""Extract Zotero citations from a Word/Google Docs document and export a BibTeX file.

Works with documents exported from Google Docs (Download as .docx) as well as
Word documents with live Zotero field codes.

Strategy (tried in order):
  1. Zotero field codes (ADDIN ZOTERO_ITEM) — available when using the Zotero
     Word plugin and citations have not been unlinked.
  2. DOI extraction from the Zotero-generated bibliography — works with Google
     Docs exports where field codes are stripped, as long as the document
     contains a Zotero bibliography with https://doi.org/ links.

Usage:
    uv run scripts/manuscript_to_bib.py <manuscript.docx> <output.bib> [--zotero-db PATH]

Zotero does not need to be running. The output .bib file can be imported into
Zotero via File → Import.
"""

from __future__ import annotations
import argparse
import json
import re
import sqlite3
import sys
from pathlib import Path


# ---------------------------------------------------------------------------
# Method 1: Zotero field codes (Word plugin / live citations)
# ---------------------------------------------------------------------------

def _extract_keys_from_field_codes(docx_path: Path) -> list[str]:
    """Extract Zotero item keys from ADDIN ZOTERO_ITEM field codes."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))
    keys: list[str] = []
    seen: set[str] = set()

    for instr in doc.element.iter(qn("w:instrText")):
        text = (instr.text or "").strip()
        if not text.startswith("ADDIN ZOTERO_ITEM"):
            continue
        json_match = re.search(r'(\{.*\})', text, re.DOTALL)
        if not json_match:
            continue
        try:
            payload = json.loads(json_match.group(1))
        except json.JSONDecodeError:
            continue
        for cite_item in payload.get("citationItems", []):
            for uri in cite_item.get("uris", []):
                m = re.search(r'/items/([A-Z0-9]{8})$', uri)
                if m:
                    key = m.group(1)
                    if key not in seen:
                        seen.add(key)
                        keys.append(key)

    return keys


# ---------------------------------------------------------------------------
# Method 2: DOI extraction from bibliography text
# ---------------------------------------------------------------------------

_DOI_RE = re.compile(r'https?://doi\.org/(10\.\S+?)(?:\s|$)', re.IGNORECASE)


def _extract_text_from_docx(docx_path: Path) -> str:
    from docx import Document
    doc = Document(str(docx_path))
    return "\n".join(p.text for p in doc.paragraphs)


def _extract_dois_from_text(text: str) -> list[str]:
    """Extract unique DOIs from document text, preserving order."""
    dois = []
    seen: set[str] = set()
    for m in _DOI_RE.finditer(text):
        doi = m.group(1).rstrip(".,;)")
        if doi not in seen:
            seen.add(doi)
            dois.append(doi)
    return dois


def _match_dois_to_zotero(dois: list[str], zotero_db: Path) -> list[str]:
    """Return Zotero item keys for matched DOIs (case-insensitive)."""
    conn = sqlite3.connect(f"file:{zotero_db}?mode=ro&immutable=1", uri=True)
    conn.row_factory = sqlite3.Row

    doi_field_id = conn.execute(
        "SELECT fieldID FROM fields WHERE fieldName = 'DOI'"
    ).fetchone()
    if not doi_field_id:
        conn.close()
        return []
    doi_field_id = doi_field_id["fieldID"]

    keys = []
    for doi in dois:
        row = conn.execute(
            """
            SELECT i.key FROM items i
            JOIN itemData d ON d.itemID = i.itemID
            JOIN itemDataValues v ON v.valueID = d.valueID
            WHERE d.fieldID = ? AND LOWER(v.value) = LOWER(?)
            AND i.itemTypeID NOT IN (14, 26)
            LIMIT 1
            """,
            (doi_field_id, doi)
        ).fetchone()
        if row:
            keys.append(row["key"])
        else:
            print(f"  Warning: DOI not found in Zotero: {doi}", file=sys.stderr)

    conn.close()
    return keys


# ---------------------------------------------------------------------------
# Fetch full metadata from Zotero for a list of keys
# ---------------------------------------------------------------------------

def _fetch_items(keys: list[str], zotero_db: Path) -> list[dict]:
    conn = sqlite3.connect(f"file:{zotero_db}?mode=ro&immutable=1", uri=True)
    conn.row_factory = sqlite3.Row

    field_rows = conn.execute(
        "SELECT fieldName, fieldID FROM fields "
        "WHERE fieldName IN ('title','abstractNote','date','DOI','url',"
        "'publicationTitle','volume','issue','pages','publisher','place',"
        "'ISBN','ISSN','edition','series')"
    ).fetchall()
    fids = {r["fieldName"]: r["fieldID"] for r in field_rows}

    author_type_id = conn.execute(
        "SELECT creatorTypeID FROM creatorTypes WHERE creatorType = 'author'"
    ).fetchone()
    author_type_id = author_type_id["creatorTypeID"] if author_type_id else 1

    items = []
    for key in keys:
        row = conn.execute(
            "SELECT itemID, itemTypeID FROM items WHERE key = ? AND itemTypeID NOT IN (14, 26)",
            (key,)
        ).fetchone()
        if row is None:
            print(f"  Warning: key {key!r} not found in Zotero library", file=sys.stderr)
            continue

        item_id = row["itemID"]

        fields: dict[str, str] = {}
        for fname, fid in fids.items():
            val = conn.execute(
                "SELECT v.value FROM itemData d JOIN itemDataValues v ON v.valueID = d.valueID "
                "WHERE d.itemID = ? AND d.fieldID = ?",
                (item_id, fid)
            ).fetchone()
            if val:
                fields[fname] = val["value"]

        authors = conn.execute(
            "SELECT c.lastName, c.firstName FROM itemCreators ic "
            "JOIN creators c ON c.creatorID = ic.creatorID "
            "WHERE ic.itemID = ? AND ic.creatorTypeID = ? "
            "ORDER BY ic.orderIndex",
            (item_id, author_type_id)
        ).fetchall()

        type_name = conn.execute(
            "SELECT typeName FROM itemTypes WHERE itemTypeID = ?", (row["itemTypeID"],)
        ).fetchone()
        type_name = type_name["typeName"] if type_name else "journalArticle"

        items.append({
            "key": key,
            "type": type_name,
            "fields": fields,
            "authors": [(a["lastName"], a["firstName"]) for a in authors],
        })

    conn.close()
    return items


# ---------------------------------------------------------------------------
# BibTeX rendering
# ---------------------------------------------------------------------------

_TYPE_MAP = {
    "journalArticle": "article",
    "book": "book",
    "bookSection": "incollection",
    "conferencePaper": "inproceedings",
    "thesis": "phdthesis",
    "report": "techreport",
    "preprint": "misc",
    "webpage": "misc",
    "document": "misc",
    "magazineArticle": "article",
    "newspaperArticle": "article",
}

_FIELD_MAP = {
    "title": "title",
    "abstractNote": "abstract",
    "publicationTitle": "journal",
    "volume": "volume",
    "issue": "number",
    "pages": "pages",
    "date": "year",
    "DOI": "doi",
    "url": "url",
    "publisher": "publisher",
    "place": "address",
    "ISBN": "isbn",
    "ISSN": "issn",
    "series": "series",
    "edition": "edition",
}


def _escape_bib(s: str) -> str:
    return s.replace("{", "\\{").replace("}", "\\}").replace("&", "\\&").replace("%", "\\%")


def _format_authors(authors: list[tuple[str, str]]) -> str:
    parts = []
    for last, first in authors:
        parts.append(f"{_escape_bib(last)}, {_escape_bib(first)}" if first else _escape_bib(last))
    return " and ".join(parts)


def _year_from_date(date_str: str) -> str:
    m = re.search(r'\b(1[89]\d{2}|2\d{3})\b', date_str)
    return m.group(1) if m else date_str[:4]


def _render_bibtex(items: list[dict]) -> str:
    lines = []
    for item in items:
        bib_type = _TYPE_MAP.get(item["type"], "misc")
        lines.append(f"@{bib_type}{{{item['key']},")
        if item["authors"]:
            lines.append(f"  author = {{{_format_authors(item['authors'])}}},")
        for zot_field, bib_field in _FIELD_MAP.items():
            val = item["fields"].get(zot_field, "")
            if not val:
                continue
            if zot_field == "date":
                val = _year_from_date(val)
                bib_field = "year"
            lines.append(f"  {bib_field} = {{{_escape_bib(val)}}},")
        lines.append("}")
        lines.append("")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument("docx", type=Path, help="Path to manuscript .docx file")
    parser.add_argument("output", type=Path, help="Output .bib file path")
    parser.add_argument(
        "--zotero-db", type=Path,
        default=Path.home() / "Zotero" / "zotero.sqlite",
        help="Path to zotero.sqlite [default: ~/Zotero/zotero.sqlite]"
    )
    args = parser.parse_args()

    if not args.docx.exists():
        sys.exit(f"Error: {args.docx} does not exist")
    if not args.zotero_db.exists():
        sys.exit(f"Error: {args.zotero_db} does not exist")
    if args.output.exists():
        sys.exit(f"Error: {args.output} already exists — delete it first or choose a different name")

    # Try method 1: field codes
    print(f"Scanning {args.docx.name} for Zotero field codes...")
    keys = _extract_keys_from_field_codes(args.docx)

    if keys:
        print(f"Found {len(keys)} cited items via field codes.")
    else:
        print("No field codes found — falling back to DOI extraction from bibliography.")
        text = _extract_text_from_docx(args.docx)
        dois = _extract_dois_from_text(text)
        if not dois:
            sys.exit(
                "No DOIs found in document. Ensure the document contains a Zotero "
                "bibliography with https://doi.org/ links, or use a document with "
                "live Zotero field codes."
            )
        print(f"Found {len(dois)} DOIs. Matching against Zotero library...")
        keys = _match_dois_to_zotero(dois, args.zotero_db)

    if not keys:
        sys.exit("No items matched in Zotero library. Nothing to export.")

    print("Fetching metadata from Zotero...")
    items = _fetch_items(keys, args.zotero_db)
    print(f"Matched {len(items)} of {len(keys)} items.")

    if not items:
        sys.exit("No items could be fetched. Nothing to export.")

    bib = _render_bibtex(items)
    args.output.write_text(bib, encoding="utf-8")
    print(f"\nWrote {args.output} ({len(items)} entries).")
    print("To import into Zotero: File → Import → BibTeX file.")


if __name__ == "__main__":
    main()
